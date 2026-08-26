"""文档摄入：加载 -> 分块 -> 嵌入 -> 写入 Milvus，并在 Postgres 记录元数据。

支持格式：txt / md / pdf / docx / html
"""
from __future__ import annotations

import hashlib
import json
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from app.config import settings
from app.db.models import Document, gen_uuid
from app.db.postgres import SessionLocal
from app.rag import image_parser, table_parser, vector_store
from app.rag.hybrid import invalidate_docs_signature

_MARKDOWN_SUFFIX = {".md", ".markdown"}
_MD_HEADERS = [("#", "H1"), ("##", "H2"), ("###", "H3")]


def _base_splitter():
    """通用分块器：中文按段落/句号/感叹号/问号切分，兼顾语义完整性。

    separators 首位加入 `=====`：兼容用等号线分章的文本（如客服知识库），
    让章节独立成块、块内语义聚焦，避免"章节说明 + 正文"混块稀释 embedding。
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    return RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=[
            "=====",
            "\n\n",
            "\n",
            "。",
            "！",
            "？",
            ". ",
            "! ",
            "? ",
            " ",
            "",
        ],
    )


def _read_text_auto(path: Path) -> str:
    """读取文本类文件，自动检测编码（chardet），检测失败回退 utf-8 容错。

    相比固定 utf-8，能正确解析 GBK 等编码的中文文档，避免乱码/丢字。
    """
    raw = path.read_bytes()
    encoding: str = "utf-8"
    try:
        import chardet

        detected = chardet.detect(raw)
        if detected:
            enc = detected.get("encoding")
            if isinstance(enc, str) and enc:
                encoding = enc
    except Exception:
        pass
    try:
        return raw.decode(encoding, errors="ignore")
    except Exception:
        return raw.decode("utf-8", errors="ignore")


class _HtmlTextExtractor(HTMLParser):
    """用 stdlib html.parser 提取 HTML 可见文本。

    - 丢弃 <script>/<style> 内容；
    - 内联标签（b/i/span/a 等）文本拼接，不随意换行；
    - 块级/换行标签（p/div/li/h1-6/table/tr/br）处换行；
    - 表格单元格用 " | " 分隔（与 docx 表一致），行末换行；
    - 列表项加 "- " 前缀。
    """

    _SKIP_TAGS = {"script", "style"}
    _BLOCK_TAGS = {
        "p", "div", "blockquote", "h1", "h2", "h3", "h4", "h5", "h6",
        "section", "article", "header", "footer", "ul", "ol", "dl", "li",
    }
    _CELL_TAGS = {"td", "th"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._lines: list[str] = []  # 普通文本行
        self._buf: list[str] = []    # 当前普通行缓冲
        self._row_parts: list[str] = []  # 当前表格行的单元格
        self._cell_buf: list[str] = []   # 当前单元格缓冲
        self._skip_depth = 0
        self._in_cell = False

    def _flush(self) -> None:
        line = "".join(self._buf).strip()
        self._buf = []
        if line:
            self._lines.append(line)

    def _flush_row(self) -> None:
        parts = [p.strip() for p in self._row_parts if p.strip()]
        if parts:
            self._lines.append(" | ".join(parts))
        self._row_parts = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in self._SKIP_TAGS:
            self._skip_depth += 1
        elif tag == "br":
            self._flush()
        elif tag == "li":
            self._flush()
            self._buf.append("- ")
        elif tag in self._CELL_TAGS:
            self._in_cell = True
            self._cell_buf = []
        elif tag == "tr":
            self._flush_row()
        elif tag in self._BLOCK_TAGS:
            self._flush()

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP_TAGS and self._skip_depth:
            self._skip_depth -= 1
            return
        if tag in self._CELL_TAGS and self._in_cell:
            self._in_cell = False
            self._row_parts.append("".join(self._cell_buf))
            self._cell_buf = []
        elif tag == "tr":
            self._flush_row()
        elif tag == "li":
            self._flush()
        elif tag in self._BLOCK_TAGS:
            self._flush()

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        if self._in_cell:
            self._cell_buf.append(data)
        else:
            self._buf.append(data)

    def text(self) -> str:
        self._flush()
        self._flush_row()
        return "\n".join(self._lines)


def _html_to_text(raw: str) -> str:
    """HTML → 纯文本（剥离标签与脚本/样式）；解析失败时回退原始文本。"""
    parser = _HtmlTextExtractor()
    try:
        parser.feed(raw)
        parser.close()
    except Exception:
        return raw
    return parser.text() or raw


def _docx_to_text(path: Path) -> str:
    """DOCX → 纯文本：段落 + 表格（每行单元格用 | 连接）。

    相比只取 paragraphs，能把表格内容也纳入知识库，不丢信息。
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _via_pdfplumber(path: Path) -> str:
    import pdfplumber

    with pdfplumber.open(str(path)) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def _via_pymupdf(path: Path) -> str:
    import pymupdf  # PyMuPDF 新名（fitz 已弃用）

    with pymupdf.open(str(path)) as doc:
        return "\n".join(page.get_text("text") for page in doc)


def _via_pypdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _pdf_to_text(path: Path) -> str:
    """PDF → 纯文本：pdfplumber→pymupdf→pypdf 逐级回退。

    相较 pypdf，pdfplumber/pymupdf 对中文/多栏/表格保留更好、控制字符更少。
    """
    for fn in (_via_pdfplumber, _via_pymupdf, _via_pypdf):
        try:
            text = fn(path)
            if text and text.strip():
                return text
        except Exception:
            continue
    return ""


def load_text(path: Path) -> str:
    """按扩展名读取文档为纯文本。"""
    suffix = path.suffix.lower()
    if suffix in (".txt", ".md", ".markdown"):
        return _read_text_auto(path)
    if suffix in (".html", ".htm"):
        return _html_to_text(_read_text_auto(path))
    if suffix == ".pdf":
        return _pdf_to_text(path)
    if suffix in (".docx",):
        return _docx_to_text(path)
    raise ValueError(f"不支持的文件类型: {suffix}")


def split_text(
    text: str, source: str, is_markdown: bool = False
) -> list[dict[str, Any]]:
    """分块：Markdown 先按标题层级切分（保留标题上下文），再递归分块。

    纯长度分块会切断标题与正文的联系；按标题分块让每个块自带语义边界，
    检索时命中块能携带所在章节标题，显著提升中文文档召回质量。
    """
    if is_markdown and text.strip().startswith("#"):
        from langchain_text_splitters import MarkdownHeaderTextSplitter

        md_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=_MD_HEADERS, strip_headers=settings.markdown_strip_headers
        )
        sections = md_splitter.split_text(text)
        base = _base_splitter()
        chunks: list[dict[str, Any]] = []
        for sec in sections:
            for piece in base.split_text(sec.page_content):
                if not piece.strip():
                    continue
                chunks.append(
                    {
                        "text": piece,
                        "metadata": {
                            **sec.metadata,
                            "source": source,
                            "chunk": len(chunks),
                        },
                    }
                )
        return chunks

    splitter = _base_splitter()
    return [
        {
            "text": chunk,
            "metadata": {"source": source, "chunk": i},
        }
        for i, chunk in enumerate(splitter.split_text(text))
    ]


def _chunk_hash(text: str) -> str:
    """文本块内容指纹（用于增量去重：内容未变的块不重复嵌入/写入）。"""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def load_document(path: Path) -> dict:
    """读取文档：返回 {"text", "tables", "images"}。

    表格与图片按开关可选解析（默认关=仅文本，行为不变）。图片 = PIL.Image 列表。
    解析失败安全降级：返回空列表，不影响文本通道。
    """
    text = load_text(path)
    suffix = path.suffix.lower()
    tables: list[dict] = []
    if settings.table_extract and suffix in (".pdf", ".docx", ".html", ".htm"):
        tables = table_parser.parse_tables(path, suffix)
    images: list = []
    if (settings.image_ocr_enabled or settings.image_vlm_enabled or settings.image_dual_channel) and suffix == ".pdf":
        images = image_parser.extract_pdf_images(path)
    return {"text": text, "tables": tables, "images": images}


def _build_table_chunks(tables: list[dict], source: str) -> list[dict]:
    """把每个表格按行列感知分块为结构化文本块（kind=table）。"""
    chunks: list[dict] = []
    mode = settings.table_to_text_mode
    max_rows = settings.table_max_rows_per_chunk
    for tbl in tables or []:
        for part in table_parser.split_table(tbl, max_rows):
            text = table_parser.table_to_text(part["headers"], part["rows"], mode)
            if not text.strip():
                continue
            chunks.append(
                {
                    "text": text,
                    "metadata": {
                        "source": source,
                        "kind": "table",
                        "columns": [str(c) for c in part["headers"]],
                    },
                }
            )
    return chunks


def _build_image_chunks(images: list, source: str) -> list[dict]:
    """对抽取的图片逐张 OCR，生成文本块（kind=ocr）；仅 OCR 开启时生效。OCR 失败/无字则跳过。"""
    if not settings.image_ocr_enabled:
        return []
    chunks: list[dict] = []
    for idx, img in enumerate(images or []):
        ocr = image_parser.ocr_text(img)
        if not ocr.strip():
            continue
        chunks.append(
            {
                "text": ocr,
                "metadata": {"source": source, "kind": "ocr", "image_index": idx},
            }
        )
    return chunks


def _build_vlm_chunks(images: list, source: str) -> list[dict]:
    """对图片逐张用 VLM 生成语义描述，生成文本块（kind=image_vlm）；失败/空则跳过。"""
    if not settings.image_vlm_enabled:
        return []
    from app.rag import vlm  # 延迟导入，避免加载时触发 openai 依赖

    chunks: list[dict] = []
    for idx, img in enumerate(images or []):
        caption = vlm.describe_image(img)
        if not caption.strip():
            continue
        chunks.append(
            {
                "text": f"[图片描述]\n{caption}",
                "metadata": {"source": source, "kind": "image_vlm", "image_index": idx},
            }
        )
    return chunks


def _write_image_vectors(images: list, source: str, user_id: str) -> int:
    """图文双通道：用多模态编码器对每张图编码 → 写 image_vectors collection。

    先删该 source 旧图片向量再写（幂等覆盖，避免 ghost）。编码 / 模型失败 → 降级为 0（不影响文本摄入）。
    """
    if not settings.image_dual_channel:
        return 0
    try:
        from app.rag.embedding import get_image_embedder

        embedder = get_image_embedder()
    except Exception:
        return 0
    try:
        vector_store.delete_image_by_source(source, user_id)
    except Exception:
        pass
    records = []
    for idx, img in enumerate(images or []):
        try:
            vec = embedder.encode_image(img)
        except Exception:
            continue
        cap = ""
        if settings.image_vlm_enabled:
            try:
                from app.rag import vlm

                cap = vlm.describe_image(img)
            except Exception:
                cap = ""
        records.append(
            {
                "user_id": user_id,
                "source": source,
                "page": 0,
                "image_index": idx,
                "caption": cap or f"[图片] {Path(source).name} 第{idx}张",
                "metadata": {"source": source, "image_index": idx, "type": "image"},
                "embedding": vec,
            }
        )
    if records:
        vector_store.add_image_vectors(records)
    return len(records)


def ingest_file(
    path: Path,
    filename: str | None = None,
    user_id: str = "default",
    progress_cb=None,
    force_reingest: bool = False,
) -> dict[str, Any]:
    """摄入单个文件到向量库（按用户隔离），返回统计信息。

    增量策略（内容指纹去重，仅作用于该用户的同名文档）：
    - 内容与已有块完全相同的块 → 跳过（不重复嵌入/写入）；
    - 新增/变化的块 → 嵌入并写入；
    - 旧文档中已消失的块 → 从 Postgres 与 Milvus 删除；
    - 整篇无变化 → 直接返回 unchanged=True（不触碰任何数据）。
    原子性：先解析成功再删旧增新，失败时旧数据不受影响。

    force_reingest：True 时跳过增量，先删除该 source 的旧块再全量重建。
    用于分块配置变化（chunk_size / strip_headers 等）的重新摄入——此时位置索引
    体系改变，增量会因 chunk_index 与保留块冲突而报 UniqueViolation。

    progress_cb(percent: int, stage: str)：摄入阶段回调（供前端展示进度）。
    """

    def _progress(percent: int, stage: str) -> None:
        if progress_cb:
            progress_cb(percent, stage)

    path = Path(path)
    source = str(path.resolve())
    filename = filename or path.name

    # 1. 解析 + 分块（失败则不触碰旧数据）
    _progress(5, "读取文件")
    doc = load_document(path)
    text = doc["text"]
    _progress(15, "分块中")
    is_markdown = path.suffix.lower() in _MARKDOWN_SUFFIX
    chunks = split_text(text, source, is_markdown=is_markdown)
    # 文档解析增强：表格结构化块 + 图片 OCR 块 + 图片 VLM 语义描述块（默认关时为空，行为不变）
    table_chunks = _build_table_chunks(doc["tables"], source)
    image_chunks = _build_image_chunks(doc["images"], source)
    vlm_chunks = _build_vlm_chunks(doc["images"], source)
    if table_chunks or image_chunks or vlm_chunks:
        _progress(22, "结构化解析")
    chunks = chunks + table_chunks + image_chunks + vlm_chunks
    if not chunks:
        # 无文本块时：仍需删旧图向量（force/换配置）或写新图向量（双通道开启）
        if force_reingest:
            vector_store.delete_image_by_source(source, user_id)
        if settings.image_dual_channel and doc["images"]:
            _write_image_vectors(doc["images"], source, user_id)
        _progress(100, "无内容可摄入")
        return {"filename": filename, "chunks": 0, "source": source}
    # 统一块序号，保证 doc_id:chunk_index 融合键唯一
    for idx, c in enumerate(chunks):
        if isinstance(c.get("metadata"), dict):
            c["metadata"]["chunk"] = idx
    new_hashes = [_chunk_hash(c["text"]) for c in chunks]
    _progress(25, "分块完成")
    # 图文双通道：图片向量写入独立 collection（默认关；放在增量判断前，保证文本 unchanged 也写）
    if settings.image_dual_channel and doc["images"]:
        _write_image_vectors(doc["images"], source, user_id)

    # 文档级内容去重：整篇内容已在库 → 跳过（跨文件同内容）
    content_hash = hashlib.sha256(
        "\n".join(c["text"] for c in chunks).encode("utf-8")
    ).hexdigest()
    if settings.doc_level_dedup:
        with SessionLocal() as db:
            dup = (
                db.query(Document.id)
                .filter(Document.user_id == user_id, Document.content_hash == content_hash)
                .first()
            )
            if dup:
                _progress(100, "内容已存在，跳过")
                return {
                    "filename": filename,
                    "chunks": len(chunks),
                    "source": source,
                    "unchanged": True,
                    "deduped": True,
                }

    if force_reingest:
        # 全量重建：先删该 source 旧块，避免分块配置变化导致的 chunk_index 冲突
        _progress(30, "全量重建")
        vector_store.delete_by_source(source, user_id)
        vector_store.delete_image_by_source(source, user_id)
        with SessionLocal() as db:
            db.query(Document).filter(
                Document.source == source, Document.user_id == user_id
            ).delete(synchronize_session=False)
            db.commit()
        stale_ids = []
        to_write = [(i, c) for i, c in enumerate(chunks)]
    else:
        # 2. 读取该用户已有同 source 块：hash -> id，用于增量对比
        existing: dict[str, str] = {}  # hash -> doc_id
        with SessionLocal() as db:
            rows = (
                db.query(Document.id, Document.text)
                .filter(Document.source == source, Document.user_id == user_id)
                .all()
            )
            for rid, rtext in rows:
                existing.setdefault(_chunk_hash(rtext), rid)

        stale_ids = [v for k, v in existing.items() if k not in set(new_hashes)]
        kept_hashes = set(existing.keys())
        to_write = [(i, c) for i, c in enumerate(chunks) if new_hashes[i] not in kept_hashes]

    # 3. 整篇无变化：直接返回，不触碰任何数据
    if not stale_ids and not to_write:
        _progress(100, "内容无变化，跳过")
        return {"filename": filename, "chunks": len(chunks), "source": source, "unchanged": True}

    # 4. 删除已消失的块（先删 Milvus 派生库，再删 Postgres 元数据，避免半库）
    if stale_ids:
        _progress(35, "清理过期分块")
        vector_store.delete_by_ids(stale_ids)
        with SessionLocal() as db:
            db.query(Document).filter(Document.id.in_(stale_ids)).delete(
                synchronize_session=False
            )
            db.commit()

    # 5. 增量：分批嵌入 + 先写 Milvus 再写 Postgres（跨库一致性 A3/B1）
    if to_write:
        from app.rag.embedding import get_embedder

        embedder = get_embedder()
        _progress(40, "生成向量嵌入")
        new_chunks = [c for _, c in to_write]
        texts = [c["text"] for c in new_chunks]
        batch = max(1, int(settings.embed_batch_size or 32))
        n = len(texts)
        vectors: list = []
        for i in range(0, n, batch):
            vectors.extend(embedder.embed_texts(texts[i : i + batch]))
            _progress(40 + int(40 * min(1, (i + batch) / max(1, n))),
                      f"嵌入 {min(i + batch, n)}/{n}")
        _progress(80, "写入向量库")
        # 先占位 doc_id，先写 Milvus（派生库），成功后再写 Postgres；PG 失败则删 Milvus 补偿
        doc_ids = [gen_uuid() for _ in new_chunks]
        vector_store.add_chunks(
            new_chunks, doc_ids=doc_ids, source=source, user_id=user_id, vectors=vectors
        )
        try:
            with SessionLocal() as db:
                for orig_i, (_, chunk) in enumerate(to_write):
                    db.add(
                        Document(
                            id=doc_ids[orig_i],
                            user_id=user_id,
                            filename=filename,
                            source=source,
                            chunk_index=orig_i,
                            text=chunk["text"],
                            metadata_json=json.dumps(chunk["metadata"], ensure_ascii=False),
                            content_hash=content_hash,
                        )
                    )
                db.commit()
        except Exception:
            # Postgres 失败 → 补偿删 Milvus 刚插的 id（避免派生库残留）
            vector_store.delete_by_ids(doc_ids)
            raise

    # 6. 失效文档集签名缓存（使 BM25 关键词通道立即包含新文档）
    invalidate_docs_signature()
    _progress(100, "完成")

    return {
        "filename": filename,
        "chunks": len(to_write),
        "source": source,
        "user_id": user_id,
        "unchanged": False,
    }


def ingest_directory(
    directory: Path, pattern: str = "*.{txt,md,pdf,docx,html}", user_id: str = "default"
) -> list[dict[str, Any]]:
    """批量摄入目录下所有支持的文件（归属指定用户的知识库）。"""
    results = []
    for file in Path(directory).rglob("*"):
        if file.suffix.lower() in (".txt", ".md", ".markdown", ".pdf", ".docx", ".html", ".htm"):
            try:
                results.append(ingest_file(file, user_id=user_id))
            except Exception as exc:
                results.append({"filename": file.name, "error": str(exc)})
    return results
