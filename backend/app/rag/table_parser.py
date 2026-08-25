"""表格解析：把 PDF / DOCX / HTML 中的表格转为结构化文本块。

目标：解决「表格乱序 / 超大表格被 chunk_size 硬截断 / 表头列语义丢失」三类问题。
产出的块走现有文本通道（同样 embedding + 混合检索），不改变向量 schema / 模型。
"""
from __future__ import annotations

from html.parser import HTMLParser

_TEXT_MODES = ("nl", "markdown")


def table_to_markdown(headers: list[str], rows: list[list[str]]) -> str:
    """二维表 → Markdown 表格（保列语义，'|' 对齐）。"""
    headers = [str(h or "") for h in headers]
    lines = [
        "| " + " | ".join(headers) + " |",
        "|" + "|".join(["---"] * max(len(headers), 1)) + "|",
    ]
    for r in rows:
        cells = [str(c or "") for c in r]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def table_to_nl(headers: list[str], rows: list[list[str]]) -> str:
    """二维表 → 自然语言（列名 + 每行 "列名=值"）。embedding 通常优于 Markdown 语法。"""
    cols = [str(h or "") for h in headers]
    lines = ["列名：" + "、".join(c for c in cols if c) if any(cols) else ""]
    for r in rows:
        cells = [str(c or "") for c in r]
        pairs = [f"{h}={c}" for h, c in zip(cols, cells) if h and c]
        lines.append("；".join(pairs) if pairs else "、".join(cells))
    return "\n".join(lines).strip() or "空表格"


def table_to_text(headers: list[str], rows: list[list[str]], mode: str = "nl") -> str:
    """按 mode 把表格序列化为文本（nl 推荐给 embedding；markdown 便于 LLM 阅读）。"""
    return table_to_markdown(headers, rows) if mode == "markdown" else table_to_nl(headers, rows)


def _make_table(data: list[list]) -> dict | None:
    """把原始二维数组（行=字符串|None）规范为 {headers, rows}。"""
    if not data:
        return None
    first = [str(c).strip() if c is not None else "" for c in data[0]]
    if not any(first) or len(data) == 1:
        # 首行无法当表头：用占位列名，全部当作数据行
        headers = [f"列{i + 1}" for i in range(len(data[0]))]
        rows = [[str(c).strip() if c is not None else "" for c in row] for row in data]
    else:
        headers = first
        rows = [[str(c).strip() if c is not None else "" for c in row] for row in data[1:]]
    return {"headers": headers, "rows": rows}


def split_table(table: dict, max_rows: int) -> list[dict]:
    """行列感知分块：大表按「表头 + 每 N 行」切块，每块自带表头上下文语义。"""
    headers = table.get("headers") or []
    rows = table.get("rows") or []
    if not rows:
        return [table]
    max_rows = max(1, int(max_rows or 10))
    return [{"headers": headers, "rows": rows[i:i + max_rows]} for i in range(0, len(rows), max_rows)]


def parse_tables(path, suffix: str = "") -> list[dict]:
    """按扩展名分发，返回 list[dict {headers, rows}]。"""
    if suffix == ".pdf":
        return _parse_pdf_tables(path)
    if suffix == ".docx":
        return _parse_docx_tables(path)
    if suffix in (".html", ".htm"):
        return _parse_html_tables(path)
    return []


# ---------------- PDF（pdfplumber）----------------
def _parse_pdf_tables(path) -> list[dict]:
    import pdfplumber

    tables: list[dict] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                for t in page.extract_tables():
                    t = [list(r) for r in t] if t else []
                    if t:
                        _add_table(tables, t)
    except Exception:
        pass
    return tables


# ---------------- DOCX（python-docx）----------------
def _parse_docx_tables(path) -> list[dict]:
    from docx import Document as DocxDocument

    tables: list[dict] = []
    try:
        doc = DocxDocument(str(path))
        for table in doc.tables:
            data = [[c.text.strip() for c in row.cells] for row in table.rows]
            if data:
                _add_table(tables, data)
    except Exception:
        pass
    return tables


# ---------------- HTML（stdlib HTMLParser，避免引入 bs4）----------------
class _HtmlTableParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.tables: list[list[list[str]]] = []
        self._cur: list[list[str]] | None = None
        self._row: list[str] | None = None
        self._cell: list[str] | None = None

    def handle_starttag(self, tag, attrs):
        if tag == "table":
            self._cur = []
        elif tag == "tr" and self._cur is not None:
            self._row = []
        elif tag in ("td", "th") and self._row is not None:
            self._cell = []

    def handle_endtag(self, tag):
        if tag in ("td", "th") and self._cell is not None:
            self._row.append("".join(self._cell).strip())
            self._cell = None
        elif tag == "tr" and self._row is not None:
            self._cur.append(self._row)
            self._row = None
        elif tag == "table" and self._cur is not None:
            self.tables.append(self._cur)
            self._cur = None

    def handle_data(self, data):
        if self._cell is not None:
            self._cell.append(data)


def _parse_html_tables(path) -> list[dict]:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    parser = _HtmlTableParser()
    parser.feed(raw)
    parser.close()
    tables: list[dict] = []
    for t in parser.tables:
        if t:
            _add_table(tables, t)
    return tables


def _add_table(tables: list[dict], data: list[list]) -> None:
    table = _make_table(data)
    if table:
        tables.append(table)
