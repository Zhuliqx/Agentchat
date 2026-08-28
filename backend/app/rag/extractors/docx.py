"""DOCX → 纯文本（段落 + 表格）。"""
from __future__ import annotations


def _docx_to_text(path) -> str:
    """DOCX → 纯文本：段落 + 表格（每行单元格用 | 连接）。

    相比只取 paragraphs，能把表格内容也纳入知识库，不丢信息。
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)
